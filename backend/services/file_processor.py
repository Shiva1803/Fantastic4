"""
FileProcessor service for extracting text content from files.

Supports PDF (via pdfplumber), plain text, DOCX, and images (via OCR).
"""

import os
from typing import Optional


class FileProcessor:
    """
    Extracts text content from uploaded files for embedding generation.
    Supports: PDF, TXT, DOCX, PNG, JPG (OCR).
    """
    
    def extract_text(self, file_path: str, file_type: Optional[str] = None) -> str:
        """
        Extract text from a file.
        
        Args:
            file_path: Absolute path to the file
            file_type: File extension (e.g. 'pdf', 'txt'). 
                       If None, inferred from file_path.
            
        Returns:
            Extracted text, or empty string on failure
        """
        if not os.path.exists(file_path):
            return ""
        
        if file_type is None:
            file_type = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf(file_path)
            elif file_type in ('txt', 'text'):
                return self._extract_txt(file_path)
            elif file_type == 'docx':
                return self._extract_docx(file_path)
            elif file_type in ('png', 'jpg', 'jpeg'):
                return self._extract_image_ocr(file_path)
            else:
                # Try reading as plain text
                return self._extract_txt(file_path)
        except Exception as e:
            print(f"[FileProcessor] Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file using pdfplumber."""
        import pdfplumber
        
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_txt(self, file_path: str) -> str:
        """Read plain text file content."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n".join(paragraphs)
        except ImportError:
            print("[FileProcessor] python-docx not installed, skipping DOCX extraction")
            return ""
    
    def _extract_image_ocr(self, file_path: str) -> str:
        """Extract text from images using OCR (pytesseract)."""
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except ImportError:
            print("[FileProcessor] pytesseract/Pillow not installed, skipping OCR")
            return ""
        except Exception as e:
            print(f"[FileProcessor] OCR failed for {file_path}: {e}")
            return ""
